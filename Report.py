from selenium import webdriver
from docx import Document
from docxcompose.composer import Composer
import time
import os

# 1. We need to gat her Variables
#urls = input("please enter the name of the text file containing the URLs: ")

number = input("please enter in the number of reports you want to generate.: ")
reportnum = int(number)
word = input("please enter a value: ")
reportstring = str(word)
url = input("please enter the name of the file that contains the urls: ")
reporturl = open(url, "r")
reportnumurl = 1000

# we need a function to open up a bunch of word documents
def ReportGen():
   for i in range (1, reportnum):
      document = Document()
      p = document.add_paragraph()
      r = p.add_run()
      saver = ('URLReport' + '%d' + '.docx') % (i)
      document.save(saver)

# Now we need to add the pictures to be added
def ReportFile():
   b = 1
   #while b < (reportnum * 100):
   for i in reporturl.readlines():
      DRIVER = 'chromedriver'
      driver = webdriver.Chrome(DRIVER)
      driver.get(i)
      time.sleep(10)
      screenshot = driver.save_screenshot('pic%d.png'  %  (b)) (width=Inches(1.0))
      driver.quit()
      #size = (0,0,680,400)
      #image = Image.open('pic%d.png')  %  (b)
      #region = image.crop(0,0,680,400)
      #region.save(('pic%d.png' % (b)), 'PNG', optimize=True, quality=95)
      b = b + 1

def ReportFinal():
    i = 1
    #while i <= (reportnum):
    for b in reporturl.readlines():
        concat = (reportstring + str(i) + '.docx')
        open(concat, 'a')
        document = Document()
        p = document.add_paragraph()
        p.add_run(str(i)+ '. ' + b)
        document.add_picture('pic%d.png' % (i))
        document.add_paragraph('-')
#document.add_paragraph('----------------------------------------------------------------------------------------------')
#document.add_paragraph(str(i)+ '. ' + b)
#document.add_picture('pic%d.png' % (i))
        document.save(concat)
        i = i + 1


def ReportCompose():
   i = 1
   for i in range(1,100):
      concat = (reportstring + str(i) + '.docx')
      concat1 = ('URLReport' + '1' + '.docx')
      os.system(("docxcompose %s %s -o %s") % (concat1, concat, concat1))
      i = i + 1
   for i in range(100,200):
      concat = (reportstring + str(i) + '.docx')
      concat2 = ('URLReport' + '2' + '.docx')
      os.system(("docxcompose %s %s -o %s") % (concat2, concat, concat2))
      i = i + 1
   for i in range(200,300):
      concat = (reportstring + str(i) + '.docx')
      concat3 = ('URLReport' + '3' + '.docx')
      os.system(("docxcompose %s %s -o %s") % (concat3, concat, concat3))
      i = i + 1
   for i in range(300,400):
      concat = (reportstring + str(i) + '.docx')
      concat3 = ('URLReport' + '4' + '.docx')
      os.system(("docxcompose %s %s -o %s") % (concat3, concat, concat3))
      i = i + 1
   for i in range(400,500):
      concat = (reportstring + str(i) + '.docx')
      concat3 = ('URLReport' + '5' + '.docx')
      os.system(("docxcompos %s %s -o %s") % (concat3, concat, concat3))
      i = i + 1
   for i in range(500,600):
      concat = (reportstring + str(i) + '.docx')
      concat3 = ('URLReport' + '6' + '.docx')
      os.system(("docxcompose %s %s -o %s") % (concat3, concat, concat3))
      i = i + 1
   for i in range(600,700):
      concat = (reportstring + str(i) + '.docx')
      concat3 = ('URLReport' + '7' + '.docx')
      os.system(("docxcompose %s %s -o %s") % (concat3, concat, concat3))
      i = i + 1
   for i in range(700,800):
      concat = (reportstring + str(i) + '.docx')
      concat3 = ('URLReport' + '8' + '.docx')
      os.system(("docxcompose %s %s -o %s") % (concat3, concat, concat3))
      i = i + 1
   for i in range(800,900):
      concat = (reportstring + str(i) + '.docx')
      concat3 = ('URLReport' + '9' + '.docx')
      os.system(("docxcompose %s %s -o %s") % (concat3, concat, concat3))
      i = i + 1
   for i in range(900,1000):
      concat = (reportstring + str(i) + '.docx')
      concat3 = ('URLReport' + '10' + '.docx')
      os.system(("docxcompose %s %s -o %s") % (concat3, concat, concat3))
      i = i + 1
   for i in range(1000,1100):
      concat = (reportstring + str(i) + '.docx')
      concat3 = ('URLReport' + '11' + '.docx')
      os.system(("docxcompose %s %s -o %s") % (concat3, concat, concat3))
      i = i + 1
   for i in range(1100,1200):
      concat = (reportstring + str(i) + '.docx')
      concat3 = ('URLReport' + '12' + '.docx')
      os.system(("docxcompose %s %s -o %s") % (concat3, concat, concat3))
      i = i + 1
   for i in range(1200,1300):
      concat = (reportstring + str(i) + '.docx')
      concat3 = ('URLReport' + '13' + '.docx')
      os.system(("docxcompose %s %s -o %s") % (concat3, concat, concat3))
      i = i + 1
   for i in range(1300,1400):
      concat = (reportstring + str(i) + '.docx')
      concat3 = ('URLReport' + '14' + '.docx')
      os.system(("docxcompose %s %s -o %s") % (concat3, concat, concat3))
      i = i + 1
   for i in range(1400,1500):
      concat = (reportstring + str(i) + '.docx')
      concat3 = ('URLReport' + '15' + '.docx')
      os.system(("docxcompose %s %s -o %s") % (concat3, concat, concat3))
      i = i + 1
   for i in range(1500,1600):
      concat = (reportstring + str(i) + '.docx')
      concat3 = ('URLReport' + '16' + '.docx')
      os.system(("docxcompose %s %s -o %s") % (concat3, concat, concat3))
      i = i + 1
   for i in range(1600,1700):
      concat = (reportstring + str(i) + '.docx')
      concat3 = ('URLReport' + '17' + '.docx')
      os.system(("docxcompose %s %s -o %s") % (concat3, concat, concat3))
      i = i + 1
   for i in range(1700,1800):
      concat = (reportstring + str(i) + '.docx')
      concat3 = ('URLReport' + '18' + '.docx')
      os.system(("docxcompose %s %s -o %s") % (concat3, concat, concat3))
      i = i + 1
   for i in range(1800,1900):
      concat = (reportstring + str(i) + '.docx')
      concat3 = ('URLReport' + '16' + '.docx')
      os.system(("docxcompose %s %s -o %s") % (concat3, concat, concat3))
      i = i + 1
   for i in range(1900,2000):
      concat = (reportstring + str(i) + '.docx')
      concat3 = ('URLReport' + '16' + '.docx')
      os.system(("docxcompose %s %s -o %s") % (concat3, concat, concat3))
      i = i + 1
   for i in range(2000,2100):
      concat = (reportstring + str(i) + '.docx')
      concat3 = ('URLReport' + '16' + '.docx')
      os.system(("docxcompose %s %s -o %s") % (concat3, concat, concat3))
      i = i + 1
   for i in range(2100,2200):
      concat = (reportstring + str(i) + '.docx')
      concat3 = ('URLReport' + '16' + '.docx')
      os.system(("docxcompose %s %s -o %s") % (concat3, concat, concat3))
      i = i + 1
   for i in range(2200,2300):
      concat = (reportstring + str(i) + '.docx')
      concat3 = ('URLReport' + '16' + '.docx')
      os.system(("docxcompose %s %s -o %s") % (concat3, concat, concat3))
      i = i + 1
   for i in range(2300,2400):
      concat = (reportstring + str(i) + '.docx')
      concat3 = ('URLReport' + '16' + '.docx')
      os.system(("docxcompose %s %s -o %s") % (concat3, concat, concat3))
      i = i + 1
   for i in range(2400,2500):
      concat = (reportstring + str(i) + '.docx')
      concat3 = ('URLReport' + '16' + '.docx')
      os.system(("docxcompose %s %s -o %s") % (concat3, concat, concat3))
      i = i + 1
   for i in range(2500,2600):
      concat = (reportstring + str(i) + '.docx')
      concat3 = ('URLReport' + '16' + '.docx')
      os.system(("docxcompose %s %s -o %s") % (concat3, concat, concat3))
      i = i + 1
   for i in range(2600,2700):
      concat = (reportstring + str(i) + '.docx')
      concat3 = ('URLReport' + '16' + '.docx')
      os.system(("docxcompose %s %s -o %s") % (concat3, concat, concat3))
      i = i + 1
   for i in range(2700,2800):
      concat = (reportstring + str(i) + '.docx')
      concat3 = ('URLReport' + '16' + '.docx')
      os.system(("docxcompose %s %s -o %s") % (concat3, concat, concat3))
      i = i + 1
   for i in range(2800,2900):
      concat = (reportstring + str(i) + '.docx')
      concat3 = ('URLReport' + '16' + '.docx')
      os.system(("docxcompose %s %s -o %s") % (concat3, concat, concat3))
      i = i + 1
   for i in range(2900,3000):
      concat = (reportstring + str(i) + '.docx')
      concat3 = ('URLReport' + '16' + '.docx')
      os.system(("docxcompose %s %s -o %s") % (concat3, concat, concat3))
      i = i + 1
   for i in range(3000,3100):
      concat = (reportstring + str(i) + '.docx')
      concat3 = ('URLReport' + '16' + '.docx')
      os.system(("docxcompose %s %s -o %s") % (concat3, concat, concat3))
      i = i + 1
   for i in range(3200, 3300):
      concat = (reportstring + str(i) + '.docx')
      concat3 = ('URLReport' + '16' + '.docx')
      os.system(("docxcompose %s %s -o %s") % (concat3, concat, concat3))
      i = i + 1
   for i in range(3300, 3400):
      concat = (reportstring + str(i) + '.docx')
      concat3 = ('URLReport' + '16' + '.docx')
      os.system(("docxcompose %s %s -o %s") % (concat3, concat, concat3))
      i = i + 1
   for i in range(3400, 3500):
      concat = (reportstring + str(i) + '.docx')
      concat3 = ('URLReport' + '16' + '.docx')
      os.system(("docxcompose %s %s -o %s") % (concat3, concat, concat3))
      i = i + 1
   for i in range(3500, 3600):
      concat = (reportstring + str(i) + '.docx')
      concat3 = ('URLReport' + '16' + '.docx')
      os.system(("docxcompose %s %s -o %s") % (concat3, concat, concat3))
      i = i + 1
   for i in range(3600, 3700):
      concat = (reportstring + str(i) + '.docx')
      concat3 = ('URLReport' + '16' + '.docx')
      os.system(("docxcompose %s %s -o %s") % (concat3, concat, concat3))
      i = i + 1
   for i in range(3700, 3800):
      concat = (reportstring + str(i) + '.docx')
      concat3 = ('URLReport' + '16' + '.docx')
      os.system(("docxcompose %s %s -o %s") % (concat3, concat, concat3))
      i = i + 1
   for i in range(3900, 4000):
      concat = (reportstring + str(i) + '.docx')
      concat3 = ('URLReport' + '16' + '.docx')
      os.system(("docxcompose %s %s -o %s") % (concat3, concat, concat3))
      i = i + 1
   for i in range(4000, 4100):
      concat = (reportstring + str(i) + '.docx')
      concat3 = ('URLReport' + '16' + '.docx')
      os.system(("docxcompose %s %s -o %s") % (concat3, concat, concat3))
      i = i + 1
   for i in range(4100, 4200):
      concat = (reportstring + str(i) + '.docx')
      concat3 = ('URLReport' + '16' + '.docx')
      os.system(("docxcompose %s %s -o %s") % (concat3, concat, concat3))
      i = i + 1
   for i in range(4200, 4300):
      concat = (reportstring + str(i) + '.docx')
      concat3 = ('URLReport' + '16' + '.docx')
      os.system(("docxcompose %s %s -o %s") % (concat3, concat, concat3))
      i = i + 1
   for i in range(4300, 4400):
      concat = (reportstring + str(i) + '.docx')
      concat3 = ('URLReport' + '16' + '.docx')
      os.system(("docxcompose %s %s -o %s") % (concat3, concat, concat3))
      i = i + 1
   for i in range(4400, 4500):
      concat = (reportstring + str(i) + '.docx')
      concat3 = ('URLReport' + '16' + '.docx')
      os.system(("docxcompose %s %s -o %s") % (concat3, concat, concat3))
      i = i + 1
   for i in range(4500, 4600):
      concat = (reportstring + str(i) + '.docx')
      concat3 = ('URLReport' + '16' + '.docx')
      os.system(("docxcompose %s %s -o %s") % (concat3, concat, concat3))
      i = i + 1
   for i in range(4600, 4700):
      concat = (reportstring + str(i) + '.docx')
      concat3 = ('URLReport' + '16' + '.docx')
      os.system(("docxcompose %s %s -o %s") % (concat3, concat, concat3))
      i = i + 1
   for i in range(4700, 4800):
      concat = (reportstring + str(i) + '.docx')
      concat3 = ('URLReport' + '16' + '.docx')
      os.system(("docxcompose %s %s -o %s") % (concat3, concat, concat3))
      i = i + 1
   for i in range(4800, 4900):
      concat = (reportstring + str(i) + '.docx')
      concat3 = ('URLReport' + '16' + '.docx')
      os.system(("docxcompose %s %s -o %s") % (concat3, concat, concat3))
      i = i + 1
   for i in range(4900, 5000):
      concat = (reportstring + str(i) + '.docx')
      concat3 = ('URLReport' + '16' + '.docx')
      os.system(("docxcompose %s %s -o %s") % (concat3, concat, concat3))
      i = i + 1
   for i in range(5000, 5100):
      concat = (reportstring + str(i) + '.docx')
      concat3 = ('URLReport' + '16' + '.docx')
      os.system(("docxcompose %s %s -o %s") % (concat3, concat, concat3))
      i = i + 1
   for i in range(5100, 5200):
      concat = (reportstring + str(i) + '.docx')
      concat3 = ('URLReport' + '16' + '.docx')
      os.system(("docxcompose %s %s -o %s") % (concat3, concat, concat3))
      i = i + 1
   for i in range(5200, 5300):
      concat = (reportstring + str(i) + '.docx')
      concat3 = ('URLReport' + '16' + '.docx')
      os.system(("docxcompose %s %s -o %s") % (concat3, concat, concat3))
      i = i + 1
   for i in range(5300, 5400):
      concat = (reportstring + str(i) + '.docx')
      concat3 = ('URLReport' + '16' + '.docx')
      os.system(("docxcompose %s %s -o %s") % (concat3, concat, concat3))
      i = i + 1
   for i in range(5400, 5500):
      concat = (reportstring + str(i) + '.docx')
      concat3 = ('URLReport' + '16' + '.docx')
      os.system(("docxcompose %s %s -o %s") % (concat3, concat, concat3))
      i = i + 1

#ReportGen()
#ReportFile()
print('error1')
ReportFinal()
print('error2')
print('error3')
ReportCompose()
